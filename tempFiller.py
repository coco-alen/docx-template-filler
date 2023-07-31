import re
import os

import docx
import openpyxl



class TempFiller():
    def __init__(self, dir, docx_fileName, excel_fileName, pattern=r"{(.*?)}"):

        if not docx_fileName.endswith(".docx"):
            print("ERROR: 文件名必须以.docx结尾 (office word file name must END with .docx)")
            input("转换失败，请检查文件名，按回车键退出 (The conversion failed, please check the file name and press Enter to exit)")
            exit()
        if not excel_fileName.endswith(".xlsx"):
            print("ERROR: 文件名必须以.xlsx结尾 (office excel file name must END with .xlsx)")
            input("转换失败，请检查文件名，按回车键退出 (The conversion failed, please check the file name and press Enter to exit)")
            exit()

        self.input_file = os.path.join(dir, docx_fileName)
        self.output_file = os.path.join(dir, docx_fileName[:-5] + "-tempFiller.docx")
        self.keyword_file = os.path.join(dir, excel_fileName)

        if not os.path.exists(self.input_file):
            print("ERROR: word源文件不存在 (word file not found)")
            input("转换失败，请检查文件是否存在，按回车键退出 (The conversion failed, please check whether the file exists, press Enter to exit)")
            exit()
        if not os.path.exists(self.keyword_file):
            print("ERROR: excel关键词文件不存在 (excel keyword file not found)")
            input("转换失败，请检查文件是否存在，按回车键退出 (The conversion failed, please check whether the file exists, press Enter to exit)")
            exit()

        self.pattern = pattern

        try:
            self.docx = docx.Document(self.input_file)
            self.keyword_sheet = openpyxl.load_workbook(self.keyword_file).active
        except:
            print("文件打开失败 (File open failed)")
            input("转换失败，请检查文件是否损坏，按回车键退出 (The conversion failed, please check whether the file is corrupted and press Enter to exit)")
            exit()
        
        print(f"导入模板文件(load template file): {self.input_file} ")
        print(f"导入关键词文件(load keyword file): {self.keyword_file} ")
        print("")

        self.keyword_dict = {}

    def load_keyword(self):
        # load keyword from excel file, and save to self.keyword_dict as dict
        for i, line in enumerate(self.keyword_sheet.rows):
            if line[0].value is not None and i != 0:
                self.keyword_dict[line[0].value] = [line[1].value, 0]  # set 0 to count

    def find_keyword(self, text):
        # find keyword in text, return a list
        match = re.findall(self.pattern, text)
        return match
    
    def replace_in_runs(self, runs, matched_keyword_list):
        # have to replace in place, or the style will be lost
        # runs breaks up {} and the keywords in it, so need to try to identify them again
        keyword_record = ""
        keyword_begin = 0
        keyword_end = 0
        record_flag = False
        for i, run in enumerate(runs):
            if "{" in run.text:
                keyword_begin = i
                keyword_record = ""
                record_flag = True

            if record_flag:
                keyword_record += run.text
                if "}" in run.text:
                    record_flag = False
                    keyword_end = i
                    keyword = self.find_keyword(keyword_record)[0]
                    
                    if keyword not in self.keyword_dict:
                        print(f"警告：关键词\"{keyword}\"不存在,跳过 (Warning: keyword \"{keyword}\" not found, skip)")
                    else:
                        self.keyword_dict[keyword][1] += 1
                        runs[keyword_begin].text = self.keyword_dict[keyword][0]
                        for j in range(keyword_begin + 1, keyword_end + 1):
                            runs[j].text = ""
            

    def replace_keyword(self):
        # Iterate over all paragraphs...
        for para in self.docx.paragraphs:
            matched_keyword_list = self.find_keyword(para.text)
            # If any keywords were matched...
            if matched_keyword_list != []:
                # Replace them in runs...
                self.replace_in_runs(para.runs, matched_keyword_list)

        # Iterate over all tables...
        for table in self.docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    matched_keyword_list = self.find_keyword(cell.text)
                    # If any keywords were matched...
                    if matched_keyword_list != []:
                        # Replace them in runs...
                        self.replace_in_runs(cell.runs, matched_keyword_list)

    def save_docx(self):
        self.docx.save(self.output_file)

    def statistic(self):
        print("")
        print("关键词使用统计结果 (keyword usage statistic result):")
        for keyword in self.keyword_dict:
            print(f"{keyword}: {self.keyword_dict[keyword][1]} ", end="")
            if self.keyword_dict[keyword][1] == 0:
                print("(未使用, not used)")
            else:
                print("")

def main(temp_filler):
    temp_filler.load_keyword()
    temp_filler.replace_keyword()
    temp_filler.save_docx()
    temp_filler.statistic()
    print("")
    print(f"转换完成 (task done)")
    print(f"生成的文件输出到(output file to): {temp_filler.output_file}")

if __name__ == "__main__":
    # dir = r"E:\Learning\tempFiller"
    # input_file = r"testFile.docx"
    # keyword_file = r"testKeyword.xlsx"

    dir = input("请输入文件夹路径 (please input dir path): ")
    input_file = input("请输入word文件名,包含.docx后缀 (Please enter the word file name including '.docx'): ")
    keyword_file = input("请输入关键词excel文件名,包含.xlsx后缀 (Please enter the keyword excel file name, including '.xlsx'): ")
    print()

    temp_filler = TempFiller(dir=dir, docx_fileName=input_file, excel_fileName=keyword_file)
    main(temp_filler)
    print("\n\n")
    input("转换结束，按回车键退出 (press enter to exit)")
    # new_doc = replace_text_in_docx(input_file)
    # new_doc.save(output_file)